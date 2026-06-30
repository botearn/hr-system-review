"""从 PDF / DOCX / HTML / 纯文本中提取简历文本。"""

from __future__ import annotations

import io
import re


def extract_pdf(raw: bytes) -> str:
    import pdfplumber

    texts: list[str] = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                texts.append(t)
    return _normalize("\n\n".join(texts))


def extract_docx(raw: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格内容也拼进来
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return _normalize("\n".join(paragraphs))


def extract_html(raw: bytes | str) -> str:
    """简单 HTML → 文本：移除 script/style，保留文本节点。"""
    from html.parser import HTMLParser

    class _Stripper(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.chunks: list[str] = []
            self.skip_depth = 0

        def handle_starttag(self, tag: str, attrs):  # type: ignore[override]
            if tag in ("script", "style", "noscript"):
                self.skip_depth += 1
            elif tag in ("br", "p", "div", "li", "tr", "h1", "h2", "h3", "h4"):
                self.chunks.append("\n")

        def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
            if tag in ("script", "style", "noscript") and self.skip_depth:
                self.skip_depth -= 1

        def handle_data(self, data: str) -> None:  # type: ignore[override]
            if not self.skip_depth:
                self.chunks.append(data)

    if isinstance(raw, bytes):
        try:
            raw = raw.decode("utf-8")
        except UnicodeDecodeError:
            raw = raw.decode("gbk", errors="ignore")
    parser = _Stripper()
    parser.feed(raw)
    return _normalize("".join(parser.chunks))


def _normalize(text: str) -> str:
    text = re.sub(r"[ \t\u3000]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract(filename: str, raw: bytes) -> str:
    """按扩展名分派。失败抛 ValueError。"""
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_pdf(raw)
    if name.endswith(".docx"):
        return extract_docx(raw)
    if name.endswith((".txt", ".md")):
        try:
            return _normalize(raw.decode("utf-8"))
        except UnicodeDecodeError:
            return _normalize(raw.decode("gbk", errors="ignore"))
    if name.endswith((".html", ".htm")):
        return extract_html(raw)
    raise ValueError(f"unsupported file type: {filename}")
